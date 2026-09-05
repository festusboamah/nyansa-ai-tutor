from io import BytesIO

from docx import Document


def build_student_note_docx(note, note_data):
    document = Document()
    document.add_heading(note.topic, level=1)
    document.add_paragraph(f"{note.subject.name} — {note.class_level}")

    for section in (note_data or {}).get("sections", []):
        document.add_heading(section.get("heading", ""), level=2)
        document.add_paragraph(section.get("text", ""))

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
