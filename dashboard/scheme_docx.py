from io import BytesIO

from docx import Document


def build_scheme_of_learning_docx(scheme, scheme_data):
    document = Document()
    document.add_heading(f"Scheme of Learning: {scheme.subject.name}", level=1)
    document.add_paragraph(f"Class: {scheme.class_level}    Term: {scheme.term}")

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Week"
    header_cells[1].text = "Topic"
    for week in (scheme_data or {}).get("weeks", []):
        cells = table.add_row().cells
        cells[0].text = str(week.get("week", ""))
        cells[1].text = str(week.get("topic", ""))

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
