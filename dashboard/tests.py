from datetime import date

from django.test import TestCase, override_settings
from unittest.mock import Mock, patch
from django.urls import reverse
from django.core.exceptions import PermissionDenied, ValidationError

from accounts.models import User
from ai_core.models import AIUsageEvent
from courses.models import Subject
from quizzes.models import Quiz, Submission
from schools.models import School, SchoolMembership

from . import ai_reports, lesson_ai
from .lesson_workflow import (
    add_lesson_comment,
    record_initial_lesson_version,
    revise_lesson_note,
    review_lesson_note,
    submit_lesson_note,
)
from .models import LessonNote, LessonNoteEvent, LessonNoteNotification, LessonNoteVersion


def _fake_response(text):
    return Mock(content=[Mock(text=text)], usage=Mock(input_tokens=10, output_tokens=20))


class LessonNoteAccessBaselineTests(TestCase):
    def setUp(self):
        self.teacher = User.objects.create_user(
            username="teacher", password="test-password", role=User.Role.TEACHER
        )
        self.other_teacher = User.objects.create_user(
            username="other-teacher", password="test-password", role=User.Role.TEACHER
        )
        self.student = User.objects.create_user(username="student", password="test-password")
        self.school = School.objects.create(name="Lesson School", slug="lesson-school")
        SchoolMembership.objects.create(
            school=self.school, user=self.teacher, role=SchoolMembership.Role.TEACHER
        )
        SchoolMembership.objects.create(
            school=self.school,
            user=self.other_teacher,
            role=SchoolMembership.Role.TEACHER,
        )
        SchoolMembership.objects.create(
            school=self.school, user=self.student, role=SchoolMembership.Role.STUDENT
        )
        self.subject = Subject.objects.create(school=self.school, name="Integrated Science")
        self.note = LessonNote.objects.create(
            teacher=self.teacher,
            subject=self.subject,
            class_level="JHS 2",
            week_ending=date(2026, 7, 31),
            strand_topic="Reproduction",
            learning_indicator="Describe the stages of reproduction.",
        )

    def test_student_cannot_open_teacher_dashboard(self):
        self.client.force_login(self.student)

        response = self.client.get(reverse("teacher_dashboard"), secure=True)

        self.assertRedirects(response, reverse("home"), fetch_redirect_response=False)

    def test_teacher_only_sees_own_lesson_note_list(self):
        self.client.force_login(self.other_teacher)

        response = self.client.get(reverse("lesson_notes_list"), secure=True)

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, self.note.strand_topic)

    def test_teacher_cannot_open_another_teachers_lesson_note(self):
        self.client.force_login(self.other_teacher)

        response = self.client.get(
            reverse("lesson_note_detail", args=[self.note.id]), secure=True
        )

        self.assertEqual(response.status_code, 404)

    def test_teacher_can_open_teacher_dashboard(self):
        self.client.force_login(self.teacher)

        response = self.client.get(reverse("teacher_dashboard"), secure=True)

        self.assertEqual(response.status_code, 200)

    @override_settings(NYANSA_DEMO_MODE=True)
    @patch("dashboard.views.generate_lesson_note", return_value=None)
    def test_demo_can_create_reviewable_lesson_without_ai_credentials(self, generate):
        self.client.force_login(self.teacher)
        response = self.client.post(
            reverse("create_lesson_note"),
            {
                "subject": self.subject.pk,
                "class_level": "JHS 2",
                "week_ending": "2026-08-07",
                "strand_topic": "Habitats",
                "content_standard": "",
                "learning_indicator": "Explain a habitat relationship.",
                "performance_indicator": "",
                "reference": "",
                "resources": "Picture cards",
                "num_days": 3,
            },
            secure=True,
        )
        note = LessonNote.objects.exclude(pk=self.note.pk).get()
        self.assertRedirects(response, reverse("lesson_note_detail", args=[note.pk]), fetch_redirect_response=False)
        self.assertEqual(note.current_version, 1)
        self.assertIn("structured Integrated Science activity", note.generated_content)

    def test_student_has_no_staff_alerts_navigation_or_page_access(self):
        self.client.force_login(self.student)

        home = self.client.get(reverse("home"), secure=True)
        notifications = self.client.get(reverse("notification_center"), secure=True)

        self.assertNotContains(home, "Alerts")
        self.assertEqual(notifications.status_code, 403)

    def test_teacher_has_staff_alerts_navigation_and_page_access(self):
        self.client.force_login(self.teacher)

        home = self.client.get(reverse("home"), secure=True)
        notifications = self.client.get(reverse("notification_center"), secure=True)

        self.assertContains(home, "Alerts")
        self.assertEqual(notifications.status_code, 200)


class LessonNoteGESFieldsTests(TestCase):
    def setUp(self):
        self.teacher = User.objects.create_user(
            username="ges-teacher", password="test-password", role=User.Role.TEACHER
        )
        self.school = School.objects.create(name="GES School", slug="ges-school")
        SchoolMembership.objects.create(
            school=self.school, user=self.teacher, role=SchoolMembership.Role.TEACHER
        )
        self.subject = Subject.objects.create(school=self.school, name="Career Technology")

    @patch("dashboard.views.generate_lesson_note")
    def test_creating_a_note_saves_the_ges_fields_and_passes_them_to_generation(self, generate):
        generate.return_value = {
            "content_standard": "Demonstrate understanding of measuring tools.",
            "learning_indicator": "B7.3.1.1.1: Classify and use measuring and marking out tools.",
            "performance_indicators": "Learners will identify measuring tools.",
            "core_competencies": "Communication and Collaboration; Critical Thinking",
            "resources": "Tape measure; ruler",
            "days": [{"day": "Monday", "starter": "s", "main": "m", "reflection": "r"}],
        }
        self.client.force_login(self.teacher)

        response = self.client.post(
            reverse("create_lesson_note"),
            {
                "subject": self.subject.pk,
                "class_level": "B7",
                "class_size": 45,
                "duration": "1 hour",
                "week_ending": "2026-08-07",
                "strand_topic": "Tools, equipment and processes",
                "sub_strand": "Measuring and marking out",
                "content_standard": "",
                "learning_indicator": "B7.3.1.1.1: Classify and use measuring and marking out tools.",
                "performance_indicator": "",
                "core_competencies": "",
                "reference": "",
                "resources": "",
                "num_days": 1,
            },
            secure=True,
        )

        note = LessonNote.objects.get(subject=self.subject)
        self.assertRedirects(response, reverse("lesson_note_detail", args=[note.pk]), fetch_redirect_response=False)
        self.assertEqual(note.class_size, 45)
        self.assertEqual(note.duration, "1 hour")
        self.assertEqual(note.sub_strand, "Measuring and marking out")
        self.assertEqual(note.core_competencies, "Communication and Collaboration; Critical Thinking")
        self.assertIn("B7.3.1.1.1", note.learning_indicator)
        generate.assert_called_once()
        self.assertEqual(generate.call_args.kwargs["sub_strand"], "Measuring and marking out")

    @patch("dashboard.views.generate_lesson_note")
    def test_downloaded_docx_contains_the_ges_fields(self, generate):
        from io import BytesIO
        from docx import Document

        generate.return_value = {
            "content_standard": "Demonstrate understanding of measuring tools.",
            "learning_indicator": "B7.3.1.1.1: Classify and use measuring tools.",
            "performance_indicators": "Learners will identify measuring tools.",
            "core_competencies": "Communication and Collaboration",
            "resources": "Tape measure",
            "days": [{"day": "Monday", "starter": "Starter text", "main": "Main text", "reflection": "Reflection text"}],
        }
        self.client.force_login(self.teacher)
        self.client.post(
            reverse("create_lesson_note"),
            {
                "subject": self.subject.pk, "class_level": "B7", "class_size": 45, "duration": "1 hour",
                "week_ending": "2026-08-07", "strand_topic": "Tools, equipment and processes",
                "sub_strand": "Measuring and marking out", "content_standard": "",
                "learning_indicator": "B7.3.1.1.1: Classify and use measuring tools.",
                "performance_indicator": "", "core_competencies": "", "reference": "", "resources": "", "num_days": 1,
            },
            secure=True,
        )
        note = LessonNote.objects.get(subject=self.subject)

        response = self.client.get(reverse("download_lesson_note_docx", args=[note.pk]), secure=True)

        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        document = Document(BytesIO(response.content))
        full_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn("Measuring and marking out", full_text)
        self.assertIn("Starter text", full_text)
        self.assertIn("Main text", full_text)


class SchemeOfLearningTests(TestCase):
    def setUp(self):
        self.teacher = User.objects.create_user(
            username="scheme-teacher", password="test-password", role=User.Role.TEACHER
        )
        self.school = School.objects.create(name="Scheme School", slug="scheme-school")
        SchoolMembership.objects.create(school=self.school, user=self.teacher, role=SchoolMembership.Role.TEACHER)
        self.subject = Subject.objects.create(school=self.school, name="Career Technology")

    @patch("dashboard.views.generate_scheme_of_learning")
    def test_creating_a_scheme_saves_the_generated_weeks(self, generate):
        from .models import SchemeOfLearning

        generate.return_value = {
            "weeks": [{"week": 1, "topic": "Personal Hygiene"}, {"week": 2, "topic": "Food Commodities"}],
        }
        self.client.force_login(self.teacher)

        response = self.client.post(
            reverse("create_scheme_of_learning"),
            {
                "subject": self.subject.pk, "class_level": "B7", "term": "Term 2",
                "num_weeks": 2, "starting_topics": "",
            },
            secure=True,
        )

        scheme = SchemeOfLearning.objects.get(subject=self.subject)
        self.assertRedirects(response, reverse("scheme_of_learning_detail", args=[scheme.pk]), fetch_redirect_response=False)
        self.assertIn("Personal Hygiene", scheme.generated_content)
        generate.assert_called_once()

    @patch("dashboard.views.generate_scheme_of_learning")
    def test_downloaded_docx_contains_the_week_topics(self, generate):
        from io import BytesIO
        from docx import Document
        from .models import SchemeOfLearning

        generate.return_value = {"weeks": [{"week": 1, "topic": "Personal Hygiene"}]}
        self.client.force_login(self.teacher)
        self.client.post(
            reverse("create_scheme_of_learning"),
            {"subject": self.subject.pk, "class_level": "B7", "term": "Term 2", "num_weeks": 1, "starting_topics": ""},
            secure=True,
        )
        scheme = SchemeOfLearning.objects.get(subject=self.subject)

        response = self.client.get(reverse("download_scheme_of_learning_docx", args=[scheme.pk]), secure=True)

        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        document = Document(BytesIO(response.content))
        full_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn("Personal Hygiene", full_text)


class StudentNoteTests(TestCase):
    def setUp(self):
        self.teacher = User.objects.create_user(
            username="student-note-teacher", password="test-password", role=User.Role.TEACHER
        )
        self.school = School.objects.create(name="Notes School", slug="notes-school")
        SchoolMembership.objects.create(school=self.school, user=self.teacher, role=SchoolMembership.Role.TEACHER)
        self.subject = Subject.objects.create(school=self.school, name="Introduction to Computers")

    @patch("dashboard.views.generate_student_notes")
    def test_creating_student_notes_saves_the_generated_sections(self, generate):
        from .models import StudentNote

        generate.return_value = {
            "sections": [{"heading": "What is a computer?", "text": "A computer is a machine that processes data."}],
        }
        self.client.force_login(self.teacher)

        response = self.client.post(
            reverse("create_student_note"),
            {"subject": self.subject.pk, "class_level": "B7", "topic": "Introduction to Computers"},
            secure=True,
        )

        note = StudentNote.objects.get(subject=self.subject)
        self.assertRedirects(response, reverse("student_note_detail", args=[note.pk]), fetch_redirect_response=False)
        self.assertIn("A computer is a machine", note.generated_content)
        generate.assert_called_once()

    @patch("dashboard.views.generate_student_notes")
    def test_downloaded_docx_contains_the_sections(self, generate):
        from io import BytesIO
        from docx import Document
        from .models import StudentNote

        generate.return_value = {
            "sections": [{"heading": "What is a computer?", "text": "A computer is a machine that processes data."}],
        }
        self.client.force_login(self.teacher)
        self.client.post(
            reverse("create_student_note"),
            {"subject": self.subject.pk, "class_level": "B7", "topic": "Introduction to Computers"},
            secure=True,
        )
        note = StudentNote.objects.get(subject=self.subject)

        response = self.client.get(reverse("download_student_note_docx", args=[note.pk]), secure=True)

        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        document = Document(BytesIO(response.content))
        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("A computer is a machine", full_text)


class CreateContentPickerTests(TestCase):
    def setUp(self):
        self.teacher = User.objects.create_user(
            username="picker-teacher", password="test-password", role=User.Role.TEACHER
        )
        self.school = School.objects.create(name="Picker School", slug="picker-school")
        SchoolMembership.objects.create(school=self.school, user=self.teacher, role=SchoolMembership.Role.TEACHER)

    def test_picker_nudges_to_create_a_subject_when_none_exist(self):
        self.client.force_login(self.teacher)
        response = self.client.get(reverse("create_content"), secure=True)
        self.assertContains(response, "Add a subject first")

    def test_picker_does_not_nudge_once_a_subject_exists(self):
        Subject.objects.create(school=self.school, name="Mathematics")
        self.client.force_login(self.teacher)
        response = self.client.get(reverse("create_content"), secure=True)
        self.assertNotContains(response, "Add a subject first")


class PersonalSchoolGenerationGateTests(TestCase):
    def setUp(self):
        from billing.models import LicensePlan, SchoolLicense

        self.school = School.objects.create(name="Solo's Classroom", slug="solo-classroom", is_personal=True)
        self.teacher = User.objects.create_user(
            username="solo-gate-teacher", password="test-password", role=User.Role.TEACHER
        )
        SchoolMembership.objects.create(school=self.school, user=self.teacher, role=SchoolMembership.Role.TEACHER)
        self.subject = Subject.objects.create(school=self.school, name="Introduction to Computers")
        plan = LicensePlan.objects.get(code="INDIVIDUAL")
        today = date(2026, 8, 1)
        self.license = SchoolLicense.objects.create(
            school=self.school, plan=plan, status=SchoolLicense.Status.TRIAL,
            current_period_start=today, current_period_end=date(2026, 8, 15),
        )

    def _use_up_free_generations(self, count):
        for _ in range(count):
            AIUsageEvent.objects.create(
                school=self.school, source=AIUsageEvent.Source.LESSON_AI, model="claude", succeeded=True,
            )

    def test_generation_allowed_helper_blocks_after_three_free_uses(self):
        from dashboard.personal_school_gate import generation_allowed
        from types import SimpleNamespace

        self._use_up_free_generations(2)
        self.assertTrue(generation_allowed(SimpleNamespace(school=self.school)))

        self._use_up_free_generations(1)
        self.assertFalse(generation_allowed(SimpleNamespace(school=self.school)))

    def test_generation_allowed_ignores_real_institutions(self):
        from dashboard.personal_school_gate import generation_allowed
        from types import SimpleNamespace

        real_school = School.objects.create(name="Real Institution", slug="real-institution")
        for _ in range(10):
            AIUsageEvent.objects.create(
                school=real_school, source=AIUsageEvent.Source.LESSON_AI, model="claude", succeeded=True,
            )

        self.assertTrue(generation_allowed(SimpleNamespace(school=real_school)))

    def test_active_license_lifts_the_gate_regardless_of_usage_count(self):
        from dashboard.personal_school_gate import generation_allowed
        from billing.models import SchoolLicense
        from types import SimpleNamespace

        self._use_up_free_generations(5)
        self.license.status = SchoolLicense.Status.ACTIVE
        self.license.save(update_fields=["status"])

        self.assertTrue(generation_allowed(SimpleNamespace(school=self.school)))

    @patch("dashboard.views.generate_lesson_note")
    def test_the_fourth_generation_attempt_is_blocked_and_sent_to_pay(self, generate):
        from billing.models import LicenseInvoice

        self._use_up_free_generations(3)
        self.client.force_login(self.teacher)

        response = self.client.post(
            reverse("create_lesson_note"),
            {
                "subject": self.subject.pk, "class_level": "B7", "week_ending": "2026-08-07",
                "strand_topic": "Introduction to Computers", "content_standard": "",
                "learning_indicator": "", "performance_indicator": "", "reference": "",
                "resources": "", "num_days": 1,
            },
            secure=True,
        )

        generate.assert_not_called()
        invoice = LicenseInvoice.objects.get(license=self.license)
        self.assertRedirects(response, reverse("billing_pay_invoice", args=[invoice.pk]), fetch_redirect_response=False)
        self.assertFalse(LessonNote.objects.filter(subject=self.subject).exists())


class LessonNoteApprovalWorkflowTests(TestCase):
    def setUp(self):
        self.school = School.objects.create(name="Approval School", slug="approval-school")
        self.teacher_user = User.objects.create_user(
            username="lesson-author", password="test-password", role=User.Role.TEACHER
        )
        self.teacher = SchoolMembership.objects.create(
            school=self.school, user=self.teacher_user, role=SchoolMembership.Role.TEACHER
        )
        self.admin_user = User.objects.create_user(username="lesson-reviewer", password="test-password")
        self.administrator = SchoolMembership.objects.create(
            school=self.school, user=self.admin_user, role=SchoolMembership.Role.SCHOOL_ADMIN
        )
        self.subject = Subject.objects.create(school=self.school, name="Science")
        self.note = LessonNote.objects.create(
            teacher=self.teacher_user,
            subject=self.subject,
            class_level="JHS 1",
            week_ending=date(2026, 8, 7),
            strand_topic="Living things",
            learning_indicator="Classify living things.",
            generated_content='{"days": [{"day": "Monday", "starter": "Observe", "main": "Classify", "reflection": "Review"}]}',
        )

    def test_complete_return_revision_approval_and_reopen_lifecycle(self):
        version = record_initial_lesson_version(note=self.note, actor=self.teacher)
        self.assertEqual(version.version_number, 1)
        submit_lesson_note(note=self.note, actor=self.teacher, message="Ready for review")
        self.note.refresh_from_db()
        self.assertEqual(self.note.status, LessonNote.Status.PENDING_REVIEW)
        self.assertEqual(
            LessonNoteNotification.objects.filter(recipient=self.administrator).count(), 1
        )

        review_lesson_note(
            note=self.note,
            actor=self.administrator,
            action="return",
            message="Add a practical activity.",
        )
        revise_lesson_note(
            note=self.note,
            actor=self.teacher,
            values={"resources": "Leaves and picture cards"},
            reason="Added requested practical resources",
        )
        self.note.refresh_from_db()
        self.assertEqual(self.note.status, LessonNote.Status.DRAFT)
        self.assertEqual(self.note.current_version, 2)

        submit_lesson_note(note=self.note, actor=self.teacher)
        review_lesson_note(note=self.note, actor=self.administrator, action="approve")
        self.note.refresh_from_db()
        self.assertEqual(self.note.status, LessonNote.Status.APPROVED)
        self.note.resources = "Silently changed"
        with self.assertRaisesMessage(ValidationError, "locked"):
            self.note.save()

        self.note.refresh_from_db()
        review_lesson_note(
            note=self.note,
            actor=self.administrator,
            action="reopen",
            message="Curriculum reference changed.",
        )
        revise_lesson_note(
            note=self.note,
            actor=self.teacher,
            values={"reference": "Updated curriculum page 14"},
            reason="Updated curriculum reference",
        )
        self.note.refresh_from_db()
        self.assertEqual(self.note.current_version, 3)
        self.assertEqual(self.note.status, LessonNote.Status.DRAFT)
        self.assertEqual(self.note.versions.count(), 3)

    def test_return_and_reopen_require_comments(self):
        record_initial_lesson_version(note=self.note, actor=self.teacher)
        submit_lesson_note(note=self.note, actor=self.teacher)
        with self.assertRaisesMessage(ValidationError, "comment is required"):
            review_lesson_note(note=self.note, actor=self.administrator, action="return")
        review_lesson_note(note=self.note, actor=self.administrator, action="approve")
        with self.assertRaisesMessage(ValidationError, "reason is required"):
            review_lesson_note(note=self.note, actor=self.administrator, action="reopen")

    def test_non_author_cannot_revise_or_submit(self):
        other_user = User.objects.create_user(
            username="other-author", password="test-password", role=User.Role.TEACHER
        )
        other = SchoolMembership.objects.create(
            school=self.school, user=other_user, role=SchoolMembership.Role.TEACHER
        )
        with self.assertRaises(PermissionDenied):
            revise_lesson_note(
                note=self.note,
                actor=other,
                values={"resources": "Other"},
                reason="Unauthorized change",
            )
        with self.assertRaises(PermissionDenied):
            submit_lesson_note(note=self.note, actor=other)

    def test_other_school_administrator_cannot_review(self):
        other_school = School.objects.create(name="Other School", slug="lesson-other-school")
        other_admin_user = User.objects.create_user(username="other-admin", password="test-password")
        other_admin = SchoolMembership.objects.create(
            school=other_school, user=other_admin_user, role=SchoolMembership.Role.SCHOOL_ADMIN
        )
        record_initial_lesson_version(note=self.note, actor=self.teacher)
        submit_lesson_note(note=self.note, actor=self.teacher)
        with self.assertRaises(PermissionDenied):
            review_lesson_note(note=self.note, actor=other_admin, action="approve")

    def test_comments_are_immutable_and_notify_counterparty(self):
        event = add_lesson_comment(
            note=self.note, actor=self.administrator, message="Clarify the reflection prompt."
        )
        self.assertEqual(event.event_type, LessonNoteEvent.EventType.COMMENT)
        self.assertTrue(
            LessonNoteNotification.objects.filter(recipient=self.teacher, lesson_note=self.note).exists()
        )
        event.message = "Changed comment"
        with self.assertRaisesMessage(ValidationError, "immutable"):
            event.save()

    def test_version_is_immutable(self):
        version = record_initial_lesson_version(note=self.note, actor=self.teacher)
        version.reason = "Changed reason"
        with self.assertRaisesMessage(ValidationError, "immutable"):
            version.save()

    def test_teacher_cannot_open_review_queue_and_admin_can(self):
        self.client.force_login(self.teacher_user)
        response = self.client.get(reverse("lesson_review_queue"), secure=True)
        self.assertRedirects(response, reverse("home"), fetch_redirect_response=False)
        self.client.force_login(self.admin_user)
        response = self.client.get(reverse("lesson_review_queue"), secure=True)
        self.assertEqual(response.status_code, 200)

    def test_author_can_submit_through_view(self):
        record_initial_lesson_version(note=self.note, actor=self.teacher)
        self.client.force_login(self.teacher_user)
        response = self.client.post(
            reverse("submit_lesson_note", args=[self.note.pk]),
            {"message": "Please review"},
            secure=True,
        )
        self.assertRedirects(response, reverse("lesson_note_detail", args=[self.note.pk]), fetch_redirect_response=False)
        self.note.refresh_from_db()
        self.assertEqual(self.note.status, LessonNote.Status.PENDING_REVIEW)


class DashboardAIHelpersTests(TestCase):
    def setUp(self):
        self.school = School.objects.create(name="AI Reports School", slug="ai-reports-school")
        self.teacher = User.objects.create_user(
            username="ai-teacher", password="test-password", role=User.Role.TEACHER
        )
        self.student = User.objects.create_user(
            username="ai-student", password="test-password", role=User.Role.STUDENT
        )
        self.subject = Subject.objects.create(school=self.school, name="Mathematics")

    @patch("ai_core.client.client")
    def test_generate_lesson_note_returns_parsed_dict_on_success(self, mock_client):
        mock_client.messages.create.return_value = _fake_response('{"content_standard": "Understand fractions"}')
        result = lesson_ai.generate_lesson_note(
            "Basic 6", "Mathematics", "2026-09-12", "Fractions", "", "Add fractions", "", "", "", 3,
        )
        self.assertEqual(result, {"content_standard": "Understand fractions"})

    @patch("ai_core.client.client")
    def test_generate_lesson_note_falls_back_to_none_on_ai_failure(self, mock_client):
        mock_client.messages.create.side_effect = RuntimeError("down")
        result = lesson_ai.generate_lesson_note(
            "Basic 6", "Mathematics", "2026-09-12", "Fractions", "", "Add fractions", "", "", "", 3,
        )
        self.assertIsNone(result)

    @patch("ai_core.client.client")
    def test_generate_student_report_returns_text_on_success(self, mock_client):
        quiz = Quiz.objects.create(subject=self.subject, teacher=self.teacher, title="Quiz 1")
        submission = Submission.objects.create(quiz=quiz, student=self.student, score=80)
        mock_client.messages.create.return_value = _fake_response("A strong performance overall.")
        result = ai_reports.generate_student_report(self.student, self.subject, [submission])
        self.assertEqual(result, "A strong performance overall.")

    @patch("ai_core.client.client")
    def test_generate_student_report_falls_back_instead_of_raising(self, mock_client):
        quiz = Quiz.objects.create(subject=self.subject, teacher=self.teacher, title="Quiz 1")
        submission = Submission.objects.create(quiz=quiz, student=self.student, score=80)
        mock_client.messages.create.side_effect = RuntimeError("down")
        result = ai_reports.generate_student_report(self.student, self.subject, [submission])
        self.assertIn("could not be generated", result)

    @patch("ai_core.client.client")
    def test_generate_student_report_logs_usage_event_for_subjects_school(self, mock_client):
        quiz = Quiz.objects.create(subject=self.subject, teacher=self.teacher, title="Quiz 1")
        submission = Submission.objects.create(quiz=quiz, student=self.student, score=80)
        mock_client.messages.create.return_value = _fake_response("A strong performance overall.")

        ai_reports.generate_student_report(self.student, self.subject, [submission])

        event = AIUsageEvent.objects.get(school=self.school)
        self.assertEqual(event.source, AIUsageEvent.Source.STUDENT_REPORTS)
