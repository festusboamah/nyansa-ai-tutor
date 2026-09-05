from io import BytesIO

from docx import Document


def _add_table(document, rows, header=None):
    table = document.add_table(rows=0, cols=2 if header is None else len(header))
    table.style = "Table Grid"
    if header:
        cells = table.add_row().cells
        for index, text in enumerate(header):
            cells[index].text = str(text)
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = "" if text is None else str(text)
    return table


def build_lesson_note_docx(note, lesson_data):
    """Builds a .docx reproducing the GES weekly-lesson-plan layout, from the
    same LessonNote fields the PDF export reads."""
    document = Document()
    document.add_heading(f"{note.subject.name}: {note.strand_topic}", level=1)

    _add_table(document, [
        ("Class", note.class_level),
        ("Class Size", note.class_size or "-"),
        ("Duration", note.duration or "-"),
        ("Subject", note.subject.name),
        ("Reference", note.reference or "-"),
        ("Week Ending", note.week_ending.strftime("%b %d, %Y")),
        ("Strand", note.strand_topic),
        ("Sub-Strand", note.sub_strand or "-"),
        ("Content Standard", note.content_standard or "-"),
        ("Indicator", note.learning_indicator),
        ("Performance Indicator(s)", note.performance_indicator or "-"),
        ("Core Competencies", note.core_competencies or "-"),
        ("Teaching/Learning Resources", note.resources or "-"),
    ])

    document.add_heading("Daily Lesson Breakdown", level=2)
    days = (lesson_data or {}).get("days", [])
    _add_table(
        document,
        [(day.get("day", ""), day.get("starter", ""), day.get("main", ""), day.get("reflection", "")) for day in days],
        header=("Day", "Starter Activity", "Main Activity", "Reflection"),
    )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
